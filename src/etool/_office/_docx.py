import os
import re
from io import BytesIO

import docx
import numpy as np
from PIL import Image


class ManagerDocx:
    @staticmethod
    def replace_words(path: str, old: str, new: str) -> str:
        """
        Replace keywords in a Word document.

        :param path: file path
        :param old: keyword to replace
        :param new: new keyword after replacement
        :return: path
        """
        if path.endswith(".docx"):
            doc = docx.Document(path)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        run.text = run.text.replace(old, new)
            doc.save(path)
        else:
            raise ValueError("Only docx file format is supported!")
        return path

    @staticmethod
    def change_forward(word_path: str, save_path: str) -> str:
        """
        Swap page width/height (landscape/portrait style flip).

        :param word_path: Word file path
        :param save_path: result file path
        :return: save_path
        """
        doc = docx.Document(word_path)
        for section in doc.sections:
            section.page_width, section.page_height = section.page_height, section.page_width
        doc.save(save_path)
        return save_path

    @staticmethod
    def get_pictures(word_path: str, result_path: str) -> str:
        """
        Extract images from a Word document and save them.

        :param word_path: Word file path
        :param result_path: image save path
        :return: image save path
        """
        if not os.path.exists(result_path):
            os.makedirs(result_path)
        doc = docx.Document(word_path)

        dict_rel = doc.part._rels
        for rel in dict_rel:
            rel = dict_rel[rel]

            if "image" in rel.target_ref:
                img_name = re.findall("/(.*)", rel.target_ref)[0]
                word_name = os.path.splitext(word_path)[0]
                if os.sep in word_name:
                    new_name = word_name.split("\\")[-1]
                else:
                    new_name = word_name.split("/")[-1]
                with Image.open(BytesIO(rel.target_part.blob)) as im:
                    w, h = im.size
                img_name = "{}-{}-{}-{}".format(new_name, w, h, img_name)

                with open(f"{result_path}/{img_name}", "wb") as f:
                    f.write(rel.target_part.blob)
            else:
                pass
        return result_path
